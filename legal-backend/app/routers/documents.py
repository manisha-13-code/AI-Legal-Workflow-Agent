from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from sqlalchemy.orm import Session
from typing import List, Optional
import os, uuid, shutil, logging
from app.database import get_db
from app import models, schemas
from app.utils.auth import get_current_user

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/documents", tags=["Documents"])
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


def extract_text_from_file(file_path: str, content_type: str, original_name: str) -> str:
    """Extract text from PDF, DOCX, DOC, TXT files."""

    # ── TXT ──────────────────────────────────────────────────────────────────
    if content_type == "text/plain":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            return f"Text file: {original_name}. Error: {e}"

    # ── PDF ──────────────────────────────────────────────────────────────────
    if content_type == "application/pdf":
        # Try PyPDF2
        try:
            import PyPDF2
            parts = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    try:
                        t = page.extract_text()
                        if t and t.strip():
                            parts.append(f"[Page {i+1}]\n{t.strip()}")
                    except Exception:
                        continue
            if parts:
                return "\n\n".join(parts)
        except ImportError:
            pass
        except Exception as e:
            logger.error(f"PyPDF2 error: {e}")

        # Try pypdf
        try:
            from pypdf import PdfReader
            parts = []
            for i, page in enumerate(PdfReader(file_path).pages):
                try:
                    t = page.extract_text()
                    if t and t.strip():
                        parts.append(f"[Page {i+1}]\n{t.strip()}")
                except Exception:
                    continue
            if parts:
                return "\n\n".join(parts)
        except ImportError:
            pass
        except Exception as e:
            logger.error(f"pypdf error: {e}")

        # Try pdfplumber
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t and t.strip():
                        parts.append(f"[Page {i+1}]\n{t.strip()}")
            if parts:
                return "\n\n".join(parts)
        except ImportError:
            pass
        except Exception as e:
            logger.error(f"pdfplumber error: {e}")

        return (
            f"PDF Document: {original_name}\n\n"
            "PDF text extraction library not installed.\n"
            "Run: pip install PyPDF2\n\n"
            "The file is uploaded but text could not be extracted for AI analysis."
        )

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            import docx
            d = docx.Document(file_path)
            lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        lines.append(row_text)
            return "\n".join(lines) if lines else f"Word Document: {original_name}"
        except ImportError:
            return (
                f"Word Document: {original_name}\n\n"
                "python-docx not installed. Run: pip install python-docx"
            )
        except Exception as e:
            return f"Word Document: {original_name}. Error: {e}"

    return f"Document: {original_name}"


@router.post("", response_model=schemas.DocumentOut, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    case_id: Optional[int] = Form(None),
    doc_type: Optional[str] = Form("Legal Document"),
    category: Optional[str] = Form("General"),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(status_code=400,
            detail=f"File type '{file.content_type}' not allowed. Use PDF, DOC, DOCX, or TXT.")

    if case_id:
        case = db.query(models.Case).filter(
            models.Case.id == case_id,
            models.Case.owner_id == current_user.id
        ).first()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

    ext = os.path.splitext(file.filename)[1]
    unique_name = f"{uuid.uuid4()}{ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_name)

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    file_size = os.path.getsize(file_path)
    content_text = extract_text_from_file(file_path, file.content_type, file.filename)
    logger.info(f"Uploaded '{file.filename}': {len(content_text)} chars extracted")

    doc = models.Document(
        filename=unique_name, original_name=file.filename,
        file_type=file.content_type, file_size=file_size,
        content_text=content_text, doc_type=doc_type,
        category=category, case_id=case_id, owner_id=current_user.id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    activity = models.Activity(
        activity_type="document_upload",
        title="Document Uploaded",
        description=f"'{file.filename}' uploaded — {len(content_text)} characters extracted",
        user_id=current_user.id, case_id=case_id,
    )
    db.add(activity)
    db.commit()
    return doc


@router.get("", response_model=List[schemas.DocumentOut])
def get_documents(
    case_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    q = db.query(models.Document).filter(models.Document.owner_id == current_user.id)
    if case_id:
        q = q.filter(models.Document.case_id == case_id)
    return q.order_by(models.Document.created_at.desc()).all()


@router.get("/{doc_id}", response_model=schemas.DocumentOut)
def get_document(doc_id: int, db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)):
    doc = db.query(models.Document).filter(
        models.Document.id == doc_id, models.Document.owner_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.delete("/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(doc_id: int, db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)):
    doc = db.query(models.Document).filter(
        models.Document.id == doc_id, models.Document.owner_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    file_path = os.path.join(UPLOAD_DIR, doc.filename)
    if os.path.exists(file_path):
        os.remove(file_path)
    db.delete(doc)
    db.commit()